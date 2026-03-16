# -*- coding: utf-8 -*-
"""Export recommendation letter to Word .docx for recommender."""
import os
import sys
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    print('Import error:', e, file=sys.stderr)
    sys.exit(1)

doc = Document()
p0 = doc.add_paragraph('研究所推薦信草稿（台塑主管・協理 AVP）')
p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
if p0.runs:
    p0.runs[0].bold = True
    p0.runs[0].font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('說明：').bold = True
p.add_run('本稿由申請人依與推薦人之共事經驗擬寫，供推薦人修改或簽名使用。請依實際申請校系置換「貴所」等用詞。推薦人為台灣化學纖維股份有限公司（台塑）協理（Assistant Vice President, AVP），撰寫時已於文中強調此身份。')
doc.add_paragraph()

doc.add_paragraph('本人現任台灣化學纖維股份有限公司（台塑）協理（Assistant Vice President, AVP）。莊偉甫先生於民國一百年十二月至一百十一年一月任職於本公司先進專案經理（Advanced Project Manager），其間於美國路易斯安那州 SunShine Project 執行階段擔任我的直屬部屬，我對其專業能力、執行力與人格特質有充分了解。在此極力推薦莊偉甫先生參加貴所研究生甄試，並就專案領導與執行力、法規遵循與風險管理兩方面，說明其於該專案之具體表現。')

doc.add_paragraph('在專案領導與執行力方面，莊偉甫在我督導下主責 SunShine Project（總投資約十二億美元之綠地石化園區）的可行性評估與 FEL-1/2/3 前端規劃。他按時完成 FEED（前端工程設計），明確界定專案範圍、成本與技術可行性，並於一○七年協助取得空污與土地使用許可，展現將複雜目標拆解為可交付成果、於跨國高壓環境下掌握時程與預算之能力，與貴所重視之策略執行及跨域協調高度契合。')

doc.add_paragraph('在法規遵循與風險管理方面，他負責 SunShine Project 之環評（EIA）與排放管制策略研擬，並與美國陸軍工兵團（U.S. Army Corps of Engineers）協調土地使用許可；其間主動辨識法規與工期風險、研擬因應方案，展現系統化之風險管理與合規意識。此外，莊偉甫於台塑任內取得甲級廢水處理專責人員（甲水）、甲級空氣污染防制專責人員（甲空）、甲級廢棄物處理技術員（甲廢）及工業工程師等證照，持續以證照與進修強化本職學能，學習態度與貴所強調之在職進修、理論實務結合一致。')

doc.add_paragraph('綜上，莊偉甫先生於 SunShine Project 期間展現清晰之邏輯分析、扎實之專案管理與執行力，以及優秀之跨部門溝通與領導特質，足見其於高壓跨國環境下達成目標並兼顧品質與合規之能力。本人認為他具備就讀貴所所需之潛力與動機，能將實務經驗與管理理論有效結合，在此鄭重推薦，相信若順利進入貴所就讀，必有更出色之表現與成就。')

doc.add_paragraph()
doc.add_paragraph('推薦人__________________')
doc.add_paragraph('職稱：台灣化學纖維股份有限公司（台塑）協理（Assistant Vice President, AVP）')
doc.add_paragraph('日期__________________')

_here = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(_here, 'RecommendationLetter_供推薦人簽名.docx')
try:
    doc.save(out_path)
    print('Saved:', out_path)
except Exception as e:
    print('Save error:', e, file=sys.stderr)
    sys.exit(1)
